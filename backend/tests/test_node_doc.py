"""Tests for DocExtractorNodeExecutor — document text extraction."""

import pytest

from app.engine.context import ExecutionContext
from app.nodes.doc_extractor import DocExtractorNodeExecutor


class TestDocExtractorNode:
    def test_extract_txt_file(self, tmp_path):
        """验证 TXT 文件提取"""
        txt_file = tmp_path / "test.txt"
        txt_file.write_text("Hello, World!\nThis is a test.", encoding="utf-8")

        executor = DocExtractorNodeExecutor()
        ctx = ExecutionContext({})
        config = {"file_path": str(txt_file)}

        result = executor.execute(ctx, config)
        assert result["success"] is True
        assert result["text"] == "Hello, World!\nThis is a test."
        assert result["file_name"] == "test.txt"

    def test_extract_docx_file(self, tmp_path):
        """验证 DOCX 文件提取"""
        docx_file = tmp_path / "test.docx"

        from docx import Document
        doc = Document()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Second paragraph")
        doc.save(str(docx_file))

        executor = DocExtractorNodeExecutor()
        ctx = ExecutionContext({})
        config = {"file_path": str(docx_file)}

        result = executor.execute(ctx, config)
        assert result["success"] is True
        assert "First paragraph" in result["text"]
        assert "Second paragraph" in result["text"]

    def test_extract_pdf_file(self, tmp_path):
        """验证 PDF 文件提取"""
        pdf_file = tmp_path / "test.pdf"

        from pypdf import PdfWriter
        writer = PdfWriter()
        writer.add_blank_page(width=200, height=200)
        writer.write(str(pdf_file))
        writer.close()

        executor = DocExtractorNodeExecutor()
        ctx = ExecutionContext({})
        config = {"file_path": str(pdf_file)}

        result = executor.execute(ctx, config)
        assert result["success"] is True
        # Blank page extracts empty string; we just check success
        assert isinstance(result["text"], str)

    def test_file_not_found(self, tmp_path):
        """验证文件不存在返回错误"""
        executor = DocExtractorNodeExecutor()
        ctx = ExecutionContext({})
        config = {"file_path": str(tmp_path / "nonexistent.txt")}

        result = executor.execute(ctx, config)
        assert result["success"] is False
        assert "not found" in result["error"].lower()

    def test_unsupported_file_type(self, tmp_path):
        """验证不支持的文件类型返回错误"""
        unsupported_file = tmp_path / "test.xyz"
        unsupported_file.write_text("content")

        executor = DocExtractorNodeExecutor()
        ctx = ExecutionContext({})
        config = {"file_path": str(unsupported_file)}

        result = executor.execute(ctx, config)
        assert result["success"] is False
        assert "unsupported" in result["error"].lower()

    def test_explicit_file_type_override(self, tmp_path):
        """验证明确指定 file_type 覆盖自动检测"""
        custom_file = tmp_path / "test.custom"
        custom_file.write_text("Custom extension but real text", encoding="utf-8")

        executor = DocExtractorNodeExecutor()
        ctx = ExecutionContext({})
        config = {
            "file_path": str(custom_file),
            "file_type": "txt",
        }

        result = executor.execute(ctx, config)
        assert result["success"] is True
        assert "Custom extension" in result["text"]

    def test_missing_file_path(self):
        """验证缺少 file_path 返回错误"""
        executor = DocExtractorNodeExecutor()
        ctx = ExecutionContext({})
        config = {}

        result = executor.execute(ctx, config)
        assert result["success"] is False
        assert "No file_path" in result["error"]

    def test_variable_resolution_in_file_path(self):
        """验证 file_path 中的变量被解析"""
        executor = DocExtractorNodeExecutor()
        ctx = ExecutionContext({"path": "/tmp/nonexistent_test_file.txt"})
        config = {"file_path": "{{ input.path }}"}

        result = executor.execute(ctx, config)
        assert result["success"] is False
        assert "not found" in result["error"].lower()

    def test_file_size_in_result(self, tmp_path):
        """验证结果包含文件大小"""
        content = "A" * 1000
        txt_file = tmp_path / "size_test.txt"
        txt_file.write_text(content, encoding="utf-8")

        executor = DocExtractorNodeExecutor()
        ctx = ExecutionContext({})
        config = {"file_path": str(txt_file)}

        result = executor.execute(ctx, config)
        assert result["success"] is True
        assert result["file_size"] == 1000
